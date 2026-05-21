#!/usr/bin/env python3
"""提取测试结果坐标并与标答对比，生成HTML报告"""
import docx
import json
import re
import os

DIR = os.path.dirname(os.path.abspath(__file__))


# ── 提取标答坐标 ──────────────────────────────────────────────
def extract_answer_coords():
    doc = docx.Document(os.path.join(DIR, "标答.docx"))
    answers = {}
    for ti, table in enumerate(doc.tables):
        text = table.rows[0].cells[0].text.strip()
        # find the JSON block
        json_match = re.search(r'\{[\s\S]*\}', text)
        if not json_match:
            continue
        try:
            data = json.loads(json_match.group())
        except json.JSONDecodeError:
            continue
        qid = data.get("question_id", str(ti + 14))
        answer_areas = data.get("answer_areas", [])
        coords_list = []
        for area in answer_areas:
            cs = area.get("coords", [])
            if cs:
                xs = [c["x"] for c in cs]
                ys = [c["y"] for c in cs]
                coords_list.append({
                    "raw": cs,
                    "bbox": [min(xs), min(ys), max(xs), max(ys)]
                })
        answers[qid] = {
            "question_type": data.get("question_type", ""),
            "coords": coords_list,
        }
    return answers


# ── 健壮的JSON解析 ──────────────────────────────────────────────
def try_parse_json(text):
    """尝试多种方式解析JSON，处理各种格式问题"""
    # fix invalid escape sequences
    def fix_escapes(s):
        # replace \_ and other invalid escapes inside strings
        result = []
        i = 0
        in_string = False
        while i < len(s):
            ch = s[i]
            if ch == '"' and (i == 0 or s[i-1] != '\\'):
                in_string = not in_string
                result.append(ch)
            elif ch == '\\' and in_string and i + 1 < len(s):
                next_ch = s[i+1]
                if next_ch in '"\\/bfnrtu':
                    result.append(ch)
                    result.append(next_ch)
                else:
                    # invalid escape, drop the backslash
                    result.append(next_ch)
                i += 2
                continue
            else:
                result.append(ch)
            i += 1
        return ''.join(result)

    # find the start of JSON content
    for i, ch in enumerate(text):
        if ch not in '[{':
            continue
        # extract a candidate block
        start_char = ch
        end_char = ']' if ch == '[' else '}'
        # find matching end bracket
        depth = 0
        in_str = False
        end_pos = -1
        for j in range(i, len(text)):
            c = text[j]
            if c == '"' and (j == 0 or text[j-1] != '\\'):
                in_str = not in_str
            if not in_str:
                if c == start_char:
                    depth += 1
                elif c == end_char:
                    depth -= 1
                    if depth == 0:
                        end_pos = j + 1
                        break
        if end_pos == -1:
            continue

        candidate = text[i:end_pos]
        fixed = fix_escapes(candidate)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            continue

    return None


# ── 提取模型输出坐标 ──────────────────────────────────────────
def extract_model_outputs():
    doc = docx.Document(os.path.join(DIR, "智能批阅_试卷结构化实验验证.docx"))
    table = doc.tables[2]  # main results table
    results = []

    for ri in range(len(table.rows)):
        row = table.rows[ri]
        model = row.cells[1].text.strip()
        prompt = row.cells[2].text.strip()
        if not model or not prompt:
            continue

        nested = row.cells[3].tables
        if not nested:
            continue

        raw_text = nested[0].rows[0].cells[0].text.strip()

        # strip language markers like "Bash\n", "JSON\n", "SQL\n```json\n", "```"
        cleaned = raw_text
        for prefix in ["Bash\n", "JSON\n", "SQL\n", "```json\n", "```json", "```"]:
            if cleaned.startswith(prefix):
                cleaned = cleaned[len(prefix):]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

        # remove <thinking>...</thinking> blocks (handle unclosed too)
        cleaned = re.sub(r'<thinking>[\s\S]*?</thinking>', '', cleaned)
        if '<thinking>' in cleaned and '</thinking>' not in cleaned:
            # unclosed thinking block - remove from <thinking> to end
            idx = cleaned.find('<thinking>')
            cleaned = cleaned[:idx].strip()

        # find the JSON block by locating matching brackets
        data = try_parse_json(cleaned)
        if data is None:
            results.append({
                "model": model,
                "prompt": prompt,
                "error": "无法解析JSON",
                "questions": {}
            })
            continue

        # extract bboxes per question
        questions = extract_bboxes(data)
        results.append({
            "model": model,
            "prompt": prompt,
            "questions": questions
        })

    return results


def extract_bboxes(data):
    """递归提取所有bbox字段，按question_id组织"""
    questions = {}

    if isinstance(data, dict):
        # could be wrapped in paper structure
        if "question_list" in data:
            for q in data["question_list"]:
                process_question(q, questions)
        elif "question_id" in data:
            process_question(data, questions)
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                if "question_id" in item:
                    process_question(item, questions)

    return questions


def process_question(q, questions):
    qid = q.get("question_id", "?")
    parts = []

    # stem bbox
    stem = q.get("stem", {})
    if isinstance(stem, dict) and "bbox" in stem:
        bbox = stem["bbox"]
        if bbox and bbox != [0, 0, 0, 0] and bbox != []:
            parts.append({
                "part": "题干",
                "bbox": bbox,
                "content": stem.get("content", "")[:60]
            })

    # sub_questions
    for sq in q.get("sub_questions", []):
        sub_id = sq.get("sub_id", "")
        sq_stem = sq.get("stem", {})
        if isinstance(sq_stem, dict) and "bbox" in sq_stem:
            bbox = sq_stem["bbox"]
            if bbox and bbox != [0, 0, 0, 0] and bbox != []:
                parts.append({
                    "part": f"小问{sub_id} 题干",
                    "bbox": bbox,
                    "content": sq_stem.get("content", "")[:60]
                })

        sol = sq.get("student_solution", {})
        if isinstance(sol, dict) and "bbox" in sol:
            bbox = sol["bbox"]
            if bbox and bbox != [0, 0, 0, 0] and bbox != []:
                parts.append({
                    "part": f"小问{sub_id} 作答",
                    "bbox": bbox,
                    "content": sol.get("content", "")[:60]
                })

    # student_solution
    sol = q.get("student_solution", {})
    if isinstance(sol, dict) and "bbox" in sol:
        bbox = sol["bbox"]
        if bbox and bbox != [0, 0, 0, 0] and bbox != []:
            parts.append({
                "part": "学生作答",
                "bbox": bbox,
                "content": sol.get("content", "")[:60]
            })

    # drawing_area
    da = q.get("drawing_area", {})
    if isinstance(da, dict) and "bbox" in da:
        bbox = da["bbox"]
        if bbox and bbox != [0, 0, 0, 0] and bbox != []:
            parts.append({
                "part": "画图区",
                "bbox": bbox,
                "content": ""
            })

    questions[qid] = parts


# ── 计算IoU ───────────────────────────────────────────────────
def compute_iou(box1, box2):
    """box: [x1, y1, x2, y2]"""
    x1 = max(box1[0], box2[0])
    y1 = max(box1[1], box2[1])
    x2 = min(box1[2], box2[2])
    y2 = min(box1[3], box2[3])

    inter = max(0, x2 - x1) * max(0, y2 - y1)
    area1 = max(0, box1[2] - box1[0]) * max(0, box1[3] - box1[1])
    area2 = max(0, box2[2] - box2[0]) * max(0, box2[3] - box2[1])
    union = area1 + area2 - inter

    if union == 0:
        return 0
    return inter / union


# ── 生成HTML ──────────────────────────────────────────────────
def generate_html(answers, results):
    # aggregate bbox per question to get overall coverage
    def aggregate_bbox(parts):
        """将所有parts的bbox合并为一个整体bbox"""
        all_coords = []
        for p in parts:
            bbox = p.get("bbox", [])
            if len(bbox) == 4 and bbox != [0, 0, 0, 0]:
                all_coords.append(bbox)
        if not all_coords:
            return None
        x1 = min(b[0] for b in all_coords)
        y1 = min(b[1] for b in all_coords)
        x2 = max(b[2] for b in all_coords)
        y2 = max(b[3] for b in all_coords)
        return [x1, y1, x2, y2]

    html = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>试卷结构化 - 坐标对比报告</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif;
       background: #f5f7fa; color: #333; padding: 20px; }
h1 { text-align: center; font-size: 22px; margin-bottom: 20px; color: #1a1a2e; }
h2 { font-size: 16px; margin: 16px 0 10px; color: #2d3436; border-left: 3px solid #6c5ce7;
     padding-left: 10px; }
table { width: 100%; border-collapse: collapse; margin-bottom: 16px; background: #fff;
        border-radius: 6px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,.08); }
th { background: #6c5ce7; color: #fff; font-weight: 500; padding: 8px 10px; text-align: left;
     font-size: 13px; }
td { padding: 7px 10px; border-bottom: 1px solid #eee; font-size: 13px; vertical-align: top; }
tr:last-child td { border-bottom: none; }
tr:hover td { background: #f8f7ff; }
.badge { display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 11px;
         font-weight: 500; }
.good { background: #d4edda; color: #155724; }
.ok { background: #fff3cd; color: #856404; }
.poor { background: #f8d7da; color: #721c24; }
.none { background: #e2e3e5; color: #383d41; }
.coords { font-family: "SF Mono", "Menlo", monospace; font-size: 12px; color: #555; }
.summary-card { background: #fff; border-radius: 8px; padding: 16px; margin-bottom: 16px;
               box-shadow: 0 1px 3px rgba(0,0,0,.08); }
.summary-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
               gap: 12px; }
.stat { text-align: center; }
.stat .num { font-size: 28px; font-weight: 700; color: #6c5ce7; }
.stat .label { font-size: 12px; color: #888; margin-top: 2px; }
.error { color: #e74c3c; font-style: italic; }
</style>
</head>
<body>
<h1>试卷结构化坐标对比报告</h1>
"""

    # summary stats
    total = len(results)
    parsed = sum(1 for r in results if "error" not in r)
    html += f"""
<div class="summary-card">
<div class="summary-grid">
  <div class="stat"><div class="num">{total}</div><div class="label">测试组合</div></div>
  <div class="stat"><div class="num">{parsed}</div><div class="label">成功解析</div></div>
  <div class="stat"><div class="num">{total - parsed}</div><div class="label">解析失败</div></div>
</div>
</div>
"""

    # 标答概览
    html += "<h2>标答坐标 (像素坐标系)</h2><table>"
    html += "<tr><th>题号</th><th>题型</th><th>坐标范围 (x1, y1, x2, y2)</th></tr>"
    for qid in sorted(answers.keys()):
        a = answers[qid]
        for c in a["coords"]:
            bbox = c["bbox"]
            html += f"""<tr><td>{qid}</td><td>{a['question_type']}</td>
            <td class="coords">[{bbox[0]}, {bbox[1]}, {bbox[2]}, {bbox[3]}]</td></tr>"""
    html += "</table>"

    # per-result comparison
    for r in results:
        model = r["model"]
        prompt = r["prompt"]
        html += f'<h2>模型: {model} | Prompt: {prompt}</h2>'

        if "error" in r:
            html += f'<p class="error">{r["error"]}</p>'
            continue

        for qid in sorted(r["questions"].keys()):
            parts = r["questions"][qid]
            agg = aggregate_bbox(parts)

            html += f'<table><tr><th colspan="5">第{qid}题</th></tr>'
            html += '<tr><th>部分</th><th>模型bbox</th><th>标答bbox</th><th>内容预览</th></tr>'

            # answer bbox for this question
            ans_bbox = None
            if qid in answers:
                for c in answers[qid]["coords"]:
                    ans_bbox = c["bbox"]

            for p in parts:
                bbox_str = str(p["bbox"]) if p["bbox"] else "—"
                ans_str = str(ans_bbox) if ans_bbox else "—"
                content = p.get("content", "")
                if len(content) > 50:
                    content = content[:50] + "..."
                html += f"""<tr>
                    <td>{p['part']}</td>
                    <td class="coords">{bbox_str}</td>
                    <td class="coords">{ans_str}</td>
                    <td>{content}</td>
                </tr>"""

            # show overall assessment
            if agg and ans_bbox:
                # note: coords may be in different scale, show both
                html += f"""<tr style="background:#f0edff">
                    <td><b>合并范围</b></td>
                    <td class="coords"><b>{agg}</b></td>
                    <td class="coords"><b>{ans_bbox}</b></td>
                    <td>模型 vs 标答</td>
                </tr>"""

            html += "</table>"

    # 横向对比表
    html += "<h2>横向对比总览</h2>"
    html += """<table>
<tr><th>模型</th><th>Prompt</th>"""

    for qid in sorted(answers.keys()):
        html += f'<th>第{qid}题</th>'
    html += '<th>坐标数量</th></tr>'

    for r in results:
        model = r["model"]
        prompt = r["prompt"]
        html += f'<tr><td>{model}</td><td>{prompt}</td>'

        total_parts = 0
        if "error" in r:
            for qid in sorted(answers.keys()):
                html += '<td class="error">解析失败</td>'
        else:
            for qid in sorted(answers.keys()):
                parts = r["questions"].get(qid, [])
                total_parts += len(parts)
                if not parts:
                    html += '<td><span class="badge none">无数据</span></td>'
                else:
                    valid = [p for p in parts if p["bbox"] and p["bbox"] != [0, 0, 0, 0]]
                    html += f'<td><span class="badge good">{len(valid)} 个坐标</span></td>'

        html += f'<td>{total_parts}</td></tr>'

    html += "</table>"

    html += """
<div style="text-align:center;color:#aaa;font-size:12px;margin-top:30px;">
  自动生成 · 坐标系差异说明：标答使用像素坐标，模型输出坐标体系可能不同
</div>
</body></html>"""

    return html


# ── Main ──────────────────────────────────────────────────────
if __name__ == "__main__":
    print("提取标答坐标...")
    answers = extract_answer_coords()
    for qid, a in answers.items():
        print(f"  题{qid}: {a}")

    print("\n提取模型输出...")
    results = extract_model_outputs()
    for r in results:
        if "error" in r:
            print(f"  {r['model']} P{r['prompt']}: ERROR - {r['error']}")
        else:
            qids = list(r["questions"].keys())
            total = sum(len(v) for v in r["questions"].values())
            print(f"  {r['model']} P{r['prompt']}: {len(qids)} questions, {total} bboxes")

    print("\n生成HTML...")
    html = generate_html(answers, results)
    out_path = os.path.join(DIR, "对比报告.html")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"已生成: {out_path}")
