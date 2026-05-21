#!/usr/bin/env python3
"""在原始试卷图片上绘制标答和模型预测的bbox，生成可视化HTML"""
import docx
import json
import re
import os
import html as html_mod

from PIL import Image, ImageDraw, ImageFont

DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(DIR, "vis_output")
os.makedirs(OUT, exist_ok=True)

SRC_IMG = os.path.join(DIR, "image.png")

PART_COLORS = {
    "题干": (220, 50, 50),
    "小问": (50, 50, 220),
    "作答": (50, 160, 50),
    "画图": (180, 100, 0),
    "学生": (160, 32, 240),
}
STD_COLOR = (0, 200, 0)  # 标答案绿色


# ── 健壮的JSON解析 ──────────────────────────────────────────────
def try_parse_json(text):
    def fix_escapes(s):
        result = []
        i = 0
        in_string = False
        while i < len(s):
            ch = s[i]
            if ch == '"' and (i == 0 or s[i-1] != '\\'):
                in_string = not in_string
                result.append(ch)
            elif ch == '\\' and in_string and i + 1 < len(s):
                next_ch = s[i+1]
                if next_ch in '"\\/bfnrtu':
                    result.append(ch)
                    result.append(next_ch)
                else:
                    result.append(next_ch)
                i += 2
                continue
            else:
                result.append(ch)
            i += 1
        return ''.join(result)

    for i, ch in enumerate(text):
        if ch not in '[{':
            continue
        start_char = ch
        end_char = ']' if ch == '[' else '}'
        depth = 0
        in_str = False
        end_pos = -1
        for j in range(i, len(text)):
            c = text[j]
            if c == '"' and (j == 0 or text[j-1] != '\\'):
                in_str = not in_str
            if not in_str:
                if c == start_char:
                    depth += 1
                elif c == end_char:
                    depth -= 1
                    if depth == 0:
                        end_pos = j + 1
                        break
        if end_pos == -1:
            continue
        candidate = text[i:end_pos]
        fixed = fix_escapes(candidate)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            continue
    return None


# ── 提取标答坐标 ──────────────────────────────────────────────
def get_std_answer_areas():
    doc = docx.Document(os.path.join(DIR, "标答.docx"))
    areas = {}
    for ti, table in enumerate(doc.tables):
        text = table.rows[0].cells[0].text.strip()
        jm = re.search(r'\{[\s\S]*\}', text)
        if not jm:
            continue
        try:
            data = json.loads(jm.group())
        except json.JSONDecodeError:
            continue
        qid = data.get("question_id", str(ti + 14))
        aas = data.get("answer_areas", [])
        if aas:
            cs = aas[0].get("coords", [])
            if cs:
                xs = [c["x"] for c in cs]
                ys = [c["y"] for c in cs]
                areas[qid] = [min(xs), min(ys), max(xs), max(ys)]
    return areas


# ── 提取所有模型输出 ──────────────────────────────────────────
def extract_all_model_outputs():
    doc = docx.Document(os.path.join(DIR, "智能批阅_试卷结构化实验验证.docx"))
    table = doc.tables[2]
    results = []

    for ri in range(len(table.rows)):
        row = table.rows[ri]
        model = row.cells[1].text.strip()
        prompt = row.cells[2].text.strip()
        if not model or not prompt:
            continue
        nested = row.cells[3].tables
        if not nested:
            continue

        raw_text = nested[0].rows[0].cells[0].text.strip()
        cleaned = raw_text
        for prefix in ["Bash\n", "JSON\n", "SQL\n", "```json\n", "```json", "```"]:
            if cleaned.startswith(prefix):
                cleaned = cleaned[len(prefix):]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()
        cleaned = re.sub(r'<thinking>[\s\S]*?</thinking>', '', cleaned)
        if '<thinking>' in cleaned and '</thinking>' not in cleaned:
            idx = cleaned.find('<thinking>')
            cleaned = cleaned[:idx].strip()

        data = try_parse_json(cleaned)
        if data is None:
            continue

        questions = extract_bboxes(data)
        results.append({"model": model, "prompt": prompt, "questions": questions})
    return results


def extract_bboxes(data):
    questions = {}
    if isinstance(data, dict):
        if "question_list" in data:
            for q in data["question_list"]:
                process_question(q, questions)
        elif "question_id" in data:
            process_question(data, questions)
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict) and "question_id" in item:
                process_question(item, questions)
    return questions


def process_question(q, questions):
    qid = q.get("question_id", "?")
    parts = []

    stem = q.get("stem", {})
    if isinstance(stem, dict) and "bbox" in stem:
        bbox = stem["bbox"]
        if bbox and bbox != [0, 0, 0, 0] and bbox != []:
            parts.append({"part": "题干", "bbox": bbox})

    for sq in q.get("sub_questions", []):
        sub_id = sq.get("sub_id", "")
        sq_stem = sq.get("stem", {})
        if isinstance(sq_stem, dict) and "bbox" in sq_stem:
            bbox = sq_stem["bbox"]
            if bbox and bbox != [0, 0, 0, 0] and bbox != []:
                parts.append({"part": f"小问{sub_id}题干", "bbox": bbox})
        sol = sq.get("student_solution", {})
        if isinstance(sol, dict) and "bbox" in sol:
            bbox = sol["bbox"]
            if bbox and bbox != [0, 0, 0, 0] and bbox != []:
                parts.append({"part": f"小问{sub_id}作答", "bbox": bbox})

    sol = q.get("student_solution", {})
    if isinstance(sol, dict) and "bbox" in sol:
        bbox = sol["bbox"]
        if bbox and bbox != [0, 0, 0, 0] and bbox != []:
            parts.append({"part": "学生作答", "bbox": bbox})

    questions[qid] = parts


# ── 绘制 ──────────────────────────────────────────────────────
def draw_overlay(std_areas, model_parts, qid, output_path, sx, sy):
    """在image.png上绘制标答区域(绿色) + 模型预测bbox(彩色)"""
    img = Image.open(SRC_IMG).copy()
    draw = ImageDraw.Draw(img)
    w, h = img.size

    font = None
    for fp in ["/System/Library/Fonts/PingFang.ttc",
               "/System/Library/Fonts/STHeiti Light.ttc"]:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 16)
                break
            except Exception:
                pass
    if font is None:
        font = ImageFont.load_default()

    # 1) 标答区域 - 绿色虚线框
    if qid in std_areas:
        ax1, ay1, ax2, ay2 = std_areas[qid]
        # draw dashed-style: 4px solid green
        for off in range(4):
            draw.rectangle([ax1-off, ay1-off, ax2+off, ay2+off], outline=STD_COLOR)
        # label
        label = "标答区域"
        bb = draw.textbbox((0, 0), label, font=font)
        tw, th = bb[2]-bb[0]+8, bb[3]-bb[1]+4
        draw.rectangle([ax1, ay1-th-4, ax1+tw, ay1-2], fill=STD_COLOR)
        draw.text((ax1+4, ay1-th-2), label, fill=(255, 255, 255), font=font)

    # 2) 模型预测 bbox - 缩放后绘制
    for p in model_parts:
        bbox = p["bbox"]
        if len(bbox) != 4:
            continue
        # 缩放: 模型坐标是1000x1000归一化
        x1 = bbox[0] * sx
        y1 = bbox[1] * sy
        x2 = bbox[2] * sx
        y2 = bbox[3] * sy
        # clamp
        lx, rx = sorted([max(0, min(x1, w)), max(0, min(x2, w))])
        ty, by = sorted([max(0, min(y1, h)), max(0, min(y2, h))])
        if rx - lx < 3 or by - ty < 3:
            continue

        # color by part type
        color = (220, 50, 50)
        for key, c in PART_COLORS.items():
            if key in p["part"]:
                color = c
                break

        for off in range(3):
            draw.rectangle([lx+off, ty+off, rx-off, by-off], outline=color)

        # label with original (normalized) coords
        text = f"{p['part']} [{bbox[0]},{bbox[1]},{bbox[2]},{bbox[3]}]"
        bb = draw.textbbox((0, 0), text, font=font)
        tw, th = bb[2]-bb[0]+8, bb[3]-bb[1]+4
        ly = max(0, ty - th - 3)
        draw.rectangle([lx, ly, lx + tw, ly + th], fill=color)
        draw.text((lx+4, ly+1), text, fill=(255, 255, 255), font=font)

    img.save(output_path, quality=92)
    return output_path


# ── 生成HTML ──────────────────────────────────────────────────
def generate_vis_html(std_areas, all_results, img_size):
    w, h = img_size
    sx, sy = w / 1000, h / 1000

    # copy source image to output
    src_rel = "source.png"
    import shutil
    shutil.copy2(SRC_IMG, os.path.join(OUT, src_rel))

    # models list
    models = []
    seen = set()
    for r in all_results:
        key = (r["model"], r["prompt"])
        if key not in seen:
            models.append(r)
            seen.add(key)

    sections = ""

    # --- 全图总览: 所有标答区域 ---
    sections += '<h2>全图标答区域总览</h2>\n'
    overview_name = "overview_std.jpeg"
    overview_path = os.path.join(OUT, overview_name)
    img = Image.open(SRC_IMG).copy()
    draw = ImageDraw.Draw(img)
    font = None
    for fp in ["/System/Library/Fonts/PingFang.ttc",
               "/System/Library/Fonts/STHeiti Light.ttc"]:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 20)
                break
            except Exception:
                pass
    if font is None:
        font = ImageFont.load_default()
    for qid in sorted(std_areas.keys()):
        ax1, ay1, ax2, ay2 = std_areas[qid]
        for off in range(5):
            draw.rectangle([ax1-off, ay1-off, ax2+off, ay2+off], outline=STD_COLOR)
        label = f"第{qid}题 标答"
        bb = draw.textbbox((0, 0), label, font=font)
        tw, th = bb[2]-bb[0]+10, bb[3]-bb[1]+6
        draw.rectangle([ax1, ay1-th-6, ax1+tw, ay1-4], fill=STD_COLOR)
        draw.text((ax1+5, ay1-th-3), label, fill=(255, 255, 255), font=font)
    img.save(overview_path, quality=92)

    sections += f'<div class="img-card"><img src="{overview_name}" style="max-width:100%;max-height:900px" /></div>\n'

    # --- 逐题逐模型 ---
    for qid in sorted(std_areas.keys()):
        sections += f'<h2>第{qid}题</h2>\n'
        sections += '<div class="tabs">\n'
        sections += f'<button class="tab-btn active" onclick="showTab(this,\'q{qid}_orig\')">原图</button>\n'

        tabs_html = f'<div class="tab-content active" id="q{qid}_orig">'
        tabs_html += f'<img src="{src_rel}" style="max-width:100%;max-height:900px" /></div>\n'

        for idx, r in enumerate(models):
            short_name = r["model"].split()[0] if " " in r["model"] else r["model"]
            tab_id = f"q{qid}_m{idx}"
            sections += f'<button class="tab-btn" onclick="showTab(this,\'{tab_id}\')">{html_mod.escape(short_name)} P{r["prompt"]}</button>\n'

            parts = r["questions"].get(qid, [])
            if not parts:
                tabs_html += f'<div class="tab-content" id="{tab_id}"><p class="no-data">无有效坐标数据</p></div>\n'
                continue

            vis_name = f"q{qid}_{r['model'].replace(' ','_').replace('.','')}_p{r['prompt']}.jpeg"
            vis_path = os.path.join(OUT, vis_name)
            draw_overlay(std_areas, parts, qid, vis_path, sx, sy)

            # info table
            rows = ""
            for p in parts:
                bbox = p["bbox"]
                scaled = [round(bbox[0]*sx), round(bbox[1]*sy), round(bbox[2]*sx), round(bbox[3]*sy)]
                rows += f'<tr><td>{html_mod.escape(p["part"])}</td><td class="coords">{bbox}</td><td class="coords">{scaled}</td></tr>\n'

            tabs_html += f'''<div class="tab-content" id="{tab_id}">
<img src="{vis_name}" style="max-width:100%;max-height:900px" />
<table class="bbox-info">
<tr><th>部分</th><th>模型坐标 (1000归一化)</th><th>像素坐标</th></tr>
{rows}
</table></div>\n'''

        sections += '</div>\n' + tabs_html

    # legend
    legend = f'<span class="legend-item"><span class="legend-color" style="background:rgb{STD_COLOR}"></span>标答区域</span>\n'
    for name, color in PART_COLORS.items():
        legend += f'<span class="legend-item"><span class="legend-color" style="background:rgb{color}"></span>{name}</span>\n'

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>试卷结构化 - bbox可视化</title>
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{ font-family:-apple-system,"PingFang SC","Microsoft YaHei",sans-serif;
       background:#f5f7fa; color:#333; padding:20px; max-width:1200px; margin:0 auto; }}
h1 {{ text-align:center; font-size:22px; margin-bottom:6px; color:#1a1a2e; }}
h2 {{ font-size:16px; margin:20px 0 8px; color:#2d3436; border-left:3px solid #6c5ce7; padding-left:10px; }}
.tabs {{ display:flex; gap:4px; margin-bottom:10px; flex-wrap:wrap; }}
.tab-btn {{ padding:6px 14px; border:1px solid #ddd; background:#fff; cursor:pointer;
           border-radius:4px; font-size:13px; }}
.tab-btn:hover {{ background:#f0edff; }}
.tab-btn.active {{ background:#6c5ce7; color:#fff; border-color:#6c5ce7; }}
.tab-content {{ display:none; background:#fff; padding:12px; border-radius:6px;
              box-shadow:0 1px 3px rgba(0,0,0,.08); margin-bottom:14px; }}
.tab-content.active {{ display:block; }}
.tab-content img {{ display:block; margin:0 auto; }}
.img-card {{ background:#fff; padding:12px; border-radius:6px; box-shadow:0 1px 3px rgba(0,0,0,.08); margin-bottom:14px; }}
.img-card img {{ display:block; margin:0 auto; }}
.bbox-info {{ width:100%; border-collapse:collapse; margin-top:10px; font-size:12px; }}
.bbox-info th {{ background:#f0edff; padding:5px 8px; text-align:left; }}
.bbox-info td {{ padding:4px 8px; border-bottom:1px solid #eee; }}
.coords {{ font-family:"SF Mono","Menlo",monospace; font-size:12px; color:#555; }}
.no-data {{ color:#999; font-style:italic; padding:20px; text-align:center; }}
.legend {{ display:flex; gap:14px; margin:8px 0 16px; flex-wrap:wrap; }}
.legend-item {{ display:flex; align-items:center; gap:4px; font-size:12px; }}
.legend-color {{ display:inline-block; width:14px; height:14px; border-radius:2px; }}
</style>
</head>
<body>
<h1>试卷结构化 bbox 可视化</h1>
<p style="text-align:center;color:#888;font-size:13px;margin-bottom:10px">
图片尺寸 {w}×{h} · 模型坐标为1000×1000归一化，已缩放至像素坐标绘制 · 绿色框=标答区域
</p>
<div class="legend">{legend}</div>
{sections}
<script>
function showTab(btn, tabId) {{
  var tabsDiv = btn.parentElement;
  var allContents = [];
  var el = tabsDiv;
  while ((el = el.nextElementSibling) && !el.classList.contains('tabs') && el.tagName !== 'H2') {{
    if (el.classList.contains('tab-content')) allContents.push(el);
  }}
  allContents.forEach(function(tc) {{ tc.classList.remove('active'); }});
  document.getElementById(tabId).classList.add('active');
  tabsDiv.querySelectorAll('.tab-btn').forEach(function(b) {{ b.classList.remove('active'); }});
  btn.classList.add('active');
}}
</script>
</body></html>"""


# ── Main ──────────────────────────────────────────────────────
if __name__ == "__main__":
    img = Image.open(SRC_IMG)
    print(f"原始图片: {img.size}")

    print("提取标答坐标...")
    std_areas = get_std_answer_areas()
    for qid, bbox in std_areas.items():
        print(f"  Q{qid}: {bbox}")

    print("\n提取模型输出...")
    all_results = extract_all_model_outputs()
    for r in all_results:
        total = sum(len(v) for v in r["questions"].values())
        print(f"  {r['model']} P{r['prompt']}: {total} bboxes")

    print("\n生成可视化...")
    html = generate_vis_html(std_areas, all_results, img.size)
    out_path = os.path.join(OUT, "bbox_visual.html")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"已生成: {out_path}")
